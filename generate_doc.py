from docx import Document
import os

def create_implementation_doc():
    doc = Document()
    
    doc.add_heading('Belvio AI Interview Agent: Overall Implementation Document', 0)
    
    # 1. Executive Summary & Objective
    doc.add_heading('1. Executive Summary & Objective', level=1)
    doc.add_heading('1.1 System Objective', level=2)
    doc.add_paragraph("The primary objective of the Belvio AI Interview Agent is to fully automate technical interviewing. It is designed to save HR professionals countless hours of manual screening while providing a standardized, unbiased, and highly scalable evaluation process for all candidates.")
    
    doc.add_heading('1.2 High-Level Overview', level=2)
    doc.add_paragraph("The system acts as an AI recruiter that can schedule interviews, join live video calls (Google Meet/Zoom), interact with candidates via voice in real-time, ask dynamic follow-up questions, and generate a comprehensive post-interview evaluation report.")
    
    # 2. High-Level Architecture
    doc.add_heading('2. High-Level Architecture', level=1)
    
    doc.add_heading('2.1 Frontend (HR Dashboard)', level=2)
    doc.add_paragraph("Built with React and Vite. It serves as the control center for HR professionals. Features include:\n- A Dashboard for scheduling new interviews by uploading resumes.\n- A Sessions view for tracking live, scheduled, and completed interviews.\n- A Reports master-detail view with role-based filtering to analyze candidate performance scores and transcripts.")
    
    doc.add_heading('2.2 Backend (Core AI Engine)', level=2)
    doc.add_paragraph("Built with FastAPI (Python), utilizing an event-driven architecture to manage concurrent live interviews without blocking. It relies on in-memory Session objects and isolated thread locks to handle state for multiple candidates simultaneously.")
    
    # 3. Database Schema & Tables
    doc.add_heading('3. Database Schema & Tables', level=1)
    doc.add_paragraph("The system is powered by Supabase (PostgreSQL). It acts as the central source of truth and features graceful fallback mechanisms to ensure live interviews do not crash if the database connection drops. The schema consists of 5 primary tables:")
    
    doc.add_heading('3.1 Candidates Table', level=2)
    doc.add_paragraph("Stores personal information about the applicant.\n- Columns: id, name, email, role.\n- Purpose: Acts as the top-level entity so that one candidate can theoretically have multiple interview sessions over time.")
    
    doc.add_heading('3.2 Sessions Table', level=2)
    doc.add_paragraph("The core tracking table for every interview.\n- Columns: id, candidate_id, bot_id, status (scheduled, in_progress, completed, stopped, incomplete), total_questions, started_at, ended_at, recording_url.\n- Context Columns: It also stores the LLM analysis of their resume (key_skills, missing_skills, jd_match_score) generated when the interview was scheduled.\n- Purpose: Maintains the real-time state of the interview and links the Recall.ai bot to the candidate.")
    
    doc.add_heading('3.3 Answers Table (Transcript)', level=2)
    doc.add_paragraph("Stores the live conversation row by row.\n- Columns: session_id, q_id, role (bot vs candidate), speaker, topic, text, category, created_at.\n- Purpose: Acts as the raw transcript buffer. When the interview ends, the evaluator fetches all rows from this table, ordered by time, to read the entire conversation.")
    
    doc.add_heading('3.4 Reports Table', level=2)
    doc.add_paragraph("Stores the final graded output.\n- Columns: session_id, plus the JSON evaluation payload (scores, feedback, pass/fail result).\n- Purpose: Upserted by the asynchronous evaluator script. This is the table that the HR Dashboard reads from when rendering the Reports UI.")
    
    doc.add_heading('3.5 Scheduled Interviews Table', level=2)
    doc.add_paragraph("Tracks future Google Meet calendar events.\n- Columns: meeting_url, scheduled_for, candidate_email, session_id, status.\n- Purpose: Allows the system to know exactly when to deploy a bot to a meeting link. When the time arrives, a bot is launched, and the scheduled_interview row is linked to the active sessions row.")

    # 4. Live Interview Flow (How the bot asks questions)
    doc.add_heading('4. Real-Time Interview Flow (Bot Mechanics)', level=1)
    
    doc.add_heading('4.1 Recall.ai Webhooks & Audio', level=2)
    doc.add_paragraph("The system uses Recall.ai to spin up virtual bots that join Google Meet/Zoom calls. Recall.ai streams live transcription webhooks to the backend, and the backend returns Text-To-Speech (TTS) audio files for the bot to speak into the meeting.")
    
    doc.add_heading('4.2 Three-State Turn Detection', level=2)
    doc.add_paragraph("To prevent the AI from interrupting candidates who pause to think, the backend uses an ultra-fast LLM call (llama-3.1-8b-instant via Groq) to evaluate live transcript snippets. It categorizes speech as:\n- COMPLETE: The candidate finished their thought.\n- INCOMPLETE: The candidate trailed off.\n- UNCERTAIN: The AI prompts the candidate to continue.")
    
    doc.add_heading('4.3 How Questions Are Asked (Gating & Follow-ups)', level=2)
    doc.add_paragraph("The bot follows a precise conversational flow:\n1. Consent: The bot starts by greeting the candidate and asking for verbal consent to be recorded.\n2. Primary Question: It reads the first technical question from the database.\n3. Active Listening: It waits for the candidate to answer, using Turn Detection to know when they finish.\n4. Gating (Evaluation): Once an answer is deemed 'COMPLETE', the AI checks it against expected 'Key Concepts'.\n5. Follow-ups: If the candidate gives a superficial answer or misses key concepts, the AI dynamically generates a specific follow-up question to probe deeper.\n6. Transition: If the answer is sufficient, it gracefully acknowledges the answer and transitions to the next topic.")
    
    # 5. Evaluation & Scoring
    doc.add_heading('5. Evaluation & Scoring Criteria', level=1)
    
    doc.add_heading('5.1 Post-Interview Processing', level=2)
    doc.add_paragraph("Upon completion of the interview, an asynchronous evaluator script is launched. Instead of relying on the lightweight live-conversation LLM, it sends the entire unedited transcript to a highly capable reasoning LLM (such as Claude 3.5 Sonnet or Gemini 1.5 Pro).")
    
    doc.add_heading('5.2 How the Bot Scores Candidates', level=2)
    doc.add_paragraph("The evaluator grades the candidate using a strict JSON schema. The scoring mechanics are as follows:\n- Question-by-Question Grading: Every single question asked during the interview is evaluated independently.\n- Numeric Score: Each question is assigned a score from 1 to 10 based on technical accuracy, depth of knowledge, and clarity of communication.\n- Reasoning: The LLM provides a detailed 2-3 sentence justification for why the specific score was awarded.\n- Total Score & Verdict: The individual scores are aggregated into a final Total Score. Based on a predefined threshold, the system automatically assigns a final verdict of 'Pass' or 'Fail'.")
    
    doc.add_heading('5.3 Result Generation', level=2)
    doc.add_paragraph("The final evaluation report is saved to the Supabase database. It is instantly reflected on the HR Dashboard under the 'Reports' tab, where recruiters can expand the candidate's profile to view the pass/fail result, the total score out of 100, and the specific feedback for every question asked.")
    
    # 6. Background Tasks & Scheduling
    doc.add_heading('6. Background Tasks & Scheduling', level=1)
    
    doc.add_heading('6.1 Scheduling & Notifications', level=2)
    doc.add_paragraph("When HR schedules an interview, the system automatically generates a Google Meet link and sends an email invitation using the SendGrid HTTP API (with a Gmail SMTP fallback).")
    
    doc.add_heading('6.2 Watchdogs', level=2)
    doc.add_paragraph("Background threads continuously monitor live sessions to enforce a 45-minute hard cap and to detect 'No Shows', automatically terminating stale bots to save resources.")
    
    file_path = os.path.join(os.getcwd(), 'Belvio_Implementation_Document_Final.docx')
    doc.save(file_path)
    print(f"Document updated and saved successfully to: {file_path}")

if __name__ == "__main__":
    create_implementation_doc()
